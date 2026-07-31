"""
Extension of Step 3 — extract the real civil unit dimensions (Annexure III
of the Word quotation template) across all priced capacity tiers, the same
way extract_excel_data.py extracted the BOQ pricing.

WHY THIS IS A SEPARATE SCRIPT FROM extract_excel_data.py:
The dimensions live inside each tier's Word (.docx) file, not the Excel
cost sheets -- they're a completely different source document with a
different structure (12 named "unit" tables, not a flat line-item list).
Keeping this separate mirrors that real difference instead of forcing two
unrelated data sources into one script.

WHAT "CIVIL UNIT" MEANS HERE:
The real template has exactly 12 tables under "ANNEXURE - III" (confirmed
by walking the document in reading order and only capturing tables between
the ANNEXURE - III and ANNEXURE - IV headings -- a naive "find all 4-row
3-column tables" approach over-matches, because the equipment spec tables
under Annexure IV use the identical shape).

Each unit table has 4 rows:
    a) Name of the unit   -> e.g. "Aeration Tank"
    b) Quantity            -> e.g. "1 No."
    c) Size/ Capacity      -> e.g. "87 m3"  (sometimes blank, or "Suitable")
    d) MOC                 -> e.g. "RCC M 20"  (material of construction)

NOT ALL UNITS SCALE WITH CAPACITY:
Some units (Sludge Holding Tank, MCC & Filter Room, Flushing Tank, Garden
Tank) have a blank or non-numeric size even in the real documents. This
script preserves that honestly -- it does NOT invent a number where the
real template has none. Those units get flagged as "fixed" rather than
interpolated later.

INPUT:  raw_excel_source/STP <capacity> cum day/*.docx
OUTPUT: data/stp_dimension_data.json

Run manually whenever the source Word templates change:
    python pricing_engine/extract_docx_dimensions.py
"""

import glob
import json
import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE_DIR = Path(__file__).resolve().parent
RAW_SOURCE_DIR = BASE_DIR / "raw_excel_source"
OUTPUT_PATH = BASE_DIR / "data" / "stp_dimension_data.json"

# Only capacities we have REAL priced BOQ data for (50-500). 550/600 are
# excluded here too, on purpose, for the same reason interpolate.py
# refuses to price them: no verified data to build from.
PRICED_CAPACITIES = [50, 100, 150, 200, 250, 300, 350, 400, 450, 500]


def find_capacity_folders() -> dict[int, Path]:
    capacity_folders: dict[int, Path] = {}
    for folder in RAW_SOURCE_DIR.iterdir():
        if not folder.is_dir():
            continue
        match = re.search(r"(\d+)", folder.name)
        if match:
            capacity_folders[int(match.group(1))] = folder
    return dict(sorted(capacity_folders.items()))


def find_docx_file(folder: Path) -> Path | None:
    """
    Find the quotation .docx inside a capacity folder. Filenames are NOT
    consistent across tiers (confirmed against the real files: "STP for
    50 cum day.docx", "STP 100 cum day.docx", "STP qt for 350 cun day.docx"
    -- note the typo "cun" -- etc), so we match by extension, not by an
    assumed exact name.
    """
    docx_files = [f for f in folder.glob("*.docx") if not f.name.startswith("~$")]
    if not docx_files:
        return None
    return docx_files[0]


def iter_block_items(document: Document):
    """
    Walk paragraphs and tables in the ACTUAL reading order they appear in
    the document. python-docx's document.tables and document.paragraphs
    give you all tables / all paragraphs separately, with no way to tell
    where a table sits relative to the surrounding headings -- which is
    exactly what we need to correctly scope "tables under Annexure III
    only", not tables anywhere shaped like a unit table.
    """
    parent_elm = document.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def extract_unit_tables(filepath: Path) -> list[Table]:
    """
    Return only the real civil unit tables (the ones under ANNEXURE - III,
    before ANNEXURE - IV begins), in document order.
    """
    document = Document(filepath)
    in_annexure_iii = False
    unit_tables = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip().upper()
            if "ANNEXURE" in text and "III" in text and "IV" not in text:
                in_annexure_iii = True
            elif "ANNEXURE" in text and "IV" in text:
                in_annexure_iii = False
        elif isinstance(block, Table):
            if in_annexure_iii and len(block.columns) == 3:
                unit_tables.append(block)

    return unit_tables


def parse_size(raw_size: str) -> dict:
    """
    Try to pull a real number out of a size string like "87 m3". If the
    field is blank or non-numeric ("Suitable"), that's preserved as-is --
    this is real data from the real document, not a gap to paper over.
    """
    raw_size = (raw_size or "").strip()
    match = re.search(r"([\d.]+)\s*m ?3", raw_size, re.IGNORECASE)
    if match:
        return {"raw": raw_size, "numeric_m3": float(match.group(1)), "is_numeric": True}
    return {"raw": raw_size, "numeric_m3": None, "is_numeric": False}


def extract_dimensions_for_tier(filepath: Path) -> list[dict]:
    units = []
    for table in extract_unit_tables(filepath):
        cells = {row.cells[0].text.strip(): row.cells[2].text.strip() for row in table.rows}
        # Row labels are literally "a)", "b)", "c)", "d)" in the real doc
        name = cells.get("a)", "")
        qty = cells.get("b)", "")
        size = cells.get("c)", "")
        moc = cells.get("d)", "")
        units.append(
            {
                "name": name,
                "quantity": qty,
                "size": parse_size(size),
                "material_of_construction": moc,
            }
        )
    return units


def main() -> None:
    capacity_folders = find_capacity_folders()
    all_tiers = {}

    for capacity in PRICED_CAPACITIES:
        folder = capacity_folders.get(capacity)
        if folder is None:
            print(f"  {capacity} cum/day: folder not found, skipping")
            continue

        docx_file = find_docx_file(folder)
        if docx_file is None:
            print(f"  {capacity} cum/day: no .docx file found, skipping")
            continue

        units = extract_dimensions_for_tier(docx_file)
        all_tiers[capacity] = units
        numeric_count = sum(1 for u in units if u["size"]["is_numeric"])
        print(f"  {capacity} cum/day: extracted {len(units)} units ({numeric_count} with numeric size) from {docx_file.name}")

    output = {
        "source_note": (
            "Extracted from ANNEXURE - III of the real company Word "
            "quotation templates. Some units have a blank or non-numeric "
            "size even in the real source documents -- this is preserved "
            "honestly (is_numeric: false), not filled in with a guess."
        ),
        "tiers": all_tiers,
    }

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_text(json.dumps(output, indent=2))
    print(f"\nWrote dimension data for {len(all_tiers)} tiers to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()