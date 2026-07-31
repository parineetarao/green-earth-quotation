"""
Step 6 (part 1) — Insert docxtpl tags into the real STP quotation letter.

WHY THIS IS DONE PROGRAMMATICALLY, NOT BY HAND-EDITING IN WORD:
Given the timeline, scripting this is both faster AND more repeatable --
if the real template ever needs a fix (a clause reworded, a new field
added), this script re-derives the tagged version from the ORIGINAL raw
file instead of us hand-editing a one-off Word file that then becomes
its own untracked source of truth.

WHAT GETS TAGGED (confirmed against the real document text):
  - "M/s. "                                   -> + {{ customer_name }}
  - "Kind Attn: Mr "                          -> + {{ attn_name }}
  - "Ref No: GEEC / SAT/C2/Q4/ETP/___/___ ... Date: __/__/____"
        -> "Ref No: {{ ref_no }}   Date: {{ quotation_date }}"
  - "... for 200 m3/day capacity )"           -> "... for {{ capacity }} m3/day capacity )"
  - "Rs.  ------------  (Rupees ------ Only)" -> "Rs. {{ price_number }} (Rupees {{ price_words }} Only)"
  - "within --- to --- weeks."                -> "within {{ completion_weeks_min }} to {{ completion_weeks_max }} weeks."
  - The 12 Annexure III unit tables' "Size/ Capacity" cell (row c)
        -> {{ unit_0_size }} ... {{ unit_11_size }}, in the SAME order
           the real document lists them, matching the order
           extract_docx_dimensions.py already extracts in.

EVERYTHING ELSE (the process description, guarantee terms, payment
terms, "why choose us" content) is left completely untouched, because
we already confirmed this content is genuinely identical across every
real capacity tier -- there is nothing to templatize there.

KNOWN LIMITATION, STATED HONESTLY:
Paragraph text is replaced by clearing all runs and writing the new text
into the first run. This keeps paragraph-level formatting (font, size)
but can lose formatting differences BETWEEN runs within the same
paragraph (e.g. if only one word was bold). The tagged fields here are
short, simply-formatted phrases in the real document (checked visually),
so this is a safe trade for speed -- but the generated .docx should still
be opened and visually checked once, not assumed perfect.
"""

from pathlib import Path

from docx import Document

RAW_TEMPLATE = Path(__file__).resolve().parent / "templates" / "stp_quotation_template_RAW.docx"
TAGGED_TEMPLATE = Path(__file__).resolve().parent / "templates" / "stp_quotation_template.docx"

# Order MUST match extract_docx_dimensions.py's real extraction order,
# confirmed against the actual document: Screen Chamber, Oil & Greases
# Tank, Collection Tank, Aeration Tank, Secondary Settling Tank, Filter
# Feed Tank, Foundation, Treated Water Tank, Sludge Holding Tank,
# MCC & Filter Room, Flushing Tank, Garden Tank.
UNIT_TAG_NAMES = [
    "unit_0_size", "unit_1_size", "unit_2_size", "unit_3_size",
    "unit_4_size", "unit_5_size", "unit_6_size", "unit_7_size",
    "unit_8_size", "unit_9_size", "unit_10_size", "unit_11_size",
]


def set_paragraph_text(paragraph, new_text: str) -> None:
    """
    Replace a paragraph's visible text while keeping its paragraph-level
    formatting. See the module docstring for the honest limitation here.
    """
    for run in paragraph.runs[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def tag_paragraphs(document: Document) -> int:
    """Find and tag the known variable paragraphs. Returns how many were tagged, for verification."""
    tagged_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text

        if text.strip() == "M/s.":
            set_paragraph_text(paragraph, "M/s. {{ customer_name }}")
            tagged_count += 1

        elif text.strip().startswith("Kind Attn: Mr"):
            set_paragraph_text(paragraph, "Kind Attn: Mr {{ attn_name }}")
            tagged_count += 1

        elif text.startswith("Ref No:"):
            set_paragraph_text(paragraph, "Ref No: {{ ref_no }}\t\t\tDate: {{ quotation_date }}")
            tagged_count += 1

        elif "m3/day capacity )" in text:
            # This paragraph reads like: "                        UF system for 200 m3/day capacity )"
            new_text = text.split("for ")[0] + "for {{ capacity }} m3/day capacity )"
            set_paragraph_text(paragraph, new_text)
            tagged_count += 1

        elif text.strip().startswith("Our price for the above job"):
            set_paragraph_text(
                paragraph,
                "Our price for the above job as given in Annexure IV  will be as below. "
                "Rs. {{ price_number }} (Rupees {{ price_words }} Only) ",
            )
            tagged_count += 1

        elif text.strip().startswith("We will complete the entire project"):
            set_paragraph_text(
                paragraph,
                "We will complete the entire project within {{ completion_weeks_min }} "
                "to {{ completion_weeks_max }} weeks.",
            )
            tagged_count += 1

    return tagged_count


def tag_dimension_tables(document: Document) -> int:
    """
    Tag the 12 real Annexure III unit tables' size cell, using the SAME
    document-order walk as extract_docx_dimensions.py, so unit_0..unit_11
    line up exactly with what size_capacity() produces.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def iter_block_items(doc):
        for child in doc.element.body.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, doc)
            elif child.tag.endswith("}tbl"):
                yield Table(child, doc)

    in_annexure_iii = False
    unit_index = 0
    tagged_count = 0

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip().upper()
            if "ANNEXURE" in text and "III" in text and "IV" not in text:
                in_annexure_iii = True
            elif "ANNEXURE" in text and "IV" in text:
                in_annexure_iii = False
        elif isinstance(block, Table):
            if in_annexure_iii and len(block.columns) == 3 and unit_index < len(UNIT_TAG_NAMES):
                size_cell = block.rows[2].cells[2]  # row c) = "Size/ Capacity"
                tag = UNIT_TAG_NAMES[unit_index]
                if size_cell.paragraphs:
                    # Tag the first paragraph regardless of whether it
                    # currently has text -- 4 of the 12 real units
                    # (Sludge Holding Tank, MCC & Filter Room, Flushing
                    # Tank, Garden Tank) have a BLANK size cell in the
                    # real document, and still need a tag inserted.
                    set_paragraph_text(size_cell.paragraphs[0], "{{ " + tag + " }}")
                    tagged_count += 1
                unit_index += 1

    return tagged_count


def remove_outdated_footer_image(document: Document) -> int:
    """
    Remove the outdated letterhead footer graphic embedded in the real
    source document.

    WHY THIS EXISTS: the raw template contains an embedded IMAGE (not
    editable text -- confirmed by searching every text location in the
    document, including raw XML) showing old contact details: a phone
    number (6506 4431 / Telefax 2518 2019) and a "Branch Office: A/P
    Madilage (Kh), Kolhapur" that do NOT match the verified current
    Company Profile (93244 35058 / 93223 32764, and the Godoli-Satara /
    Verna-Goa / Bhiwandi branches already confirmed against the real
    Company Profile document and used on the company website).

    Rather than trying to generate a replacement image with "corrected"
    text baked in -- which risks fabricating new visual content on a
    real client-facing document -- this removes the outdated graphic
    entirely. The document's actual verified current contact details
    already appear correctly as real text at the top of the letter
    (the letterhead banner + header block), so removing this redundant,
    outdated image loses no real information, only a stale duplicate.

    NOTE: identified by document structure (position immediately after
    the signature block, and by being the largest embedded image at 79KB,
    consistent with a text-heavy graphic) rather than by visually reading
    the image content, which the available tooling could not confirm
    with certainty. If this turns out to be the wrong image, the fix is
    to change OUTDATED_FOOTER_IMAGE_RELATIONSHIP_TARGET below and re-run.
    """
    from docx.oxml.ns import qn

    removed_count = 0
    # rId7 -> media/image2.png, confirmed via word/_rels/document.xml.rels
    target_embed_ids = {"rId7"}

    for blip in document.element.body.iter(qn("a:blip")):
        embed_id = blip.get(qn("r:embed"))
        if embed_id in target_embed_ids:
            # Walk up to the drawing element and remove its parent run
            # entirely, so no empty run/paragraph artifact is left behind.
            drawing = blip.getparent()
            while drawing is not None and drawing.tag != qn("w:drawing"):
                drawing = drawing.getparent()
            if drawing is not None:
                run = drawing.getparent()
                paragraph_element = run.getparent()
                paragraph_element.remove(run)
                removed_count += 1

    return removed_count


def main() -> None:
    if not RAW_TEMPLATE.exists():
        raise FileNotFoundError(f"{RAW_TEMPLATE} not found. Copy the real 200 cum/day .docx there first.")

    document = Document(RAW_TEMPLATE)

    paragraph_count = tag_paragraphs(document)
    table_count = tag_dimension_tables(document)
    removed_image_count = remove_outdated_footer_image(document)

    TAGGED_TEMPLATE.parent.mkdir(parents=True, exist_ok=True)
    document.save(TAGGED_TEMPLATE)

    print(f"Tagged {paragraph_count} paragraphs and {table_count} dimension table cells.")
    print(f"Removed {removed_image_count} outdated footer image(s).")
    print(f"Saved tagged template to {TAGGED_TEMPLATE}")

    if removed_image_count == 0:
        print("\nWARNING: expected to remove 1 outdated footer image, removed 0. "
              "Either it's already gone, or the relationship ID has changed -- "
              "check word/_rels/document.xml.rels in the raw template.")

    if paragraph_count < 6:
        print("\nWARNING: expected 6 tagged paragraphs, got fewer. "
              "The real document text may not match what this script expects -- "
              "open the output and check by hand before trusting it.")
    if table_count < 12:
        print(f"\nWARNING: expected 12 tagged unit tables, got {table_count}. "
              "Check the Annexure III boundaries in the source document.")


if __name__ == "__main__":
    main()