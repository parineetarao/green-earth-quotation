"""
Step 7 (part 2) — Generate the Annexure IV Bill of Quantities as its own
document, matching the real Excel sheet's layout (Sr. No., Equipment,
Specification, Qty, Unit Rate, Amount, then a Total row).

WHY SEPARATE FROM generate_docx.py:
Confirmed against the real source documents -- this was never embedded
inside the Word letter, it's a companion priced sheet. Keeping it as its
own generator mirrors that real structure instead of forcing one
artificial combined template.
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "pricing_engine"))
from interpolate import price_capacity, load_tier_data  # noqa: E402


def generate_boq_docx(capacity: float, customer_name: str, output_path: Path) -> None:
    pricing_result = price_capacity(capacity, load_tier_data())

    if not pricing_result.in_verified_range:
        raise ValueError(f"Cannot generate a BOQ for {capacity} cum/day: {pricing_result.note}")

    document = Document()

    title = document.add_paragraph()
    title_run = title.add_run(f"ANNEXURE - IV : PRICE BID for {capacity} m3/day STP")
    title_run.bold = True
    title_run.font.size = Pt(14)

    document.add_paragraph(f"M/s. {customer_name}")

    if not pricing_result.is_exact_known_tier:
        note = document.add_paragraph()
        note_run = note.add_run(
            f"Note: pricing interpolated between verified {pricing_result.based_on_tiers[0]} "
            f"and {pricing_result.based_on_tiers[1]} cum/day tiers. Review before sending."
        )
        note_run.italic = True

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, heading in zip(header_cells, ["Sr. No.", "Equipment", "Specification", "Qty", "Unit Rate", "Amount"]):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for item in pricing_result.line_items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.sr_no)
        row_cells[1].text = item.name
        row_cells[2].text = item.specification or ""
        row_cells[3].text = str(item.qty)
        row_cells[4].text = f"{item.unit_rate:,.2f}"
        row_cells[5].text = f"{item.amount:,.2f}"

    total_row = table.add_row().cells
    total_row[4].text = "TOTAL"
    total_row[4].paragraphs[0].runs[0].bold = True
    total_row[5].text = f"{pricing_result.total:,.2f}"
    total_row[5].paragraphs[0].runs[0].bold = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


if __name__ == "__main__":
    output_path = (
        Path(__file__).resolve().parent.parent / "storage" / "generated_quotations" / "test_boq_200.docx"
    )
    generate_boq_docx(capacity=200, customer_name="Test Industries Pvt. Ltd.", output_path=output_path)
    print(f"Generated: {output_path}")
